from docx import Document
from docx.shared import Pt

path = 'SW-ML-17 Concept Paper.docx'
doc = Document(path)

heading_order = [
    ('INTRODUCTION', 'I. INTRODUCTION'),
    ('BACKGROUND AND LITERATURE REVIEW', 'II. BACKGROUND AND LITERATURE REVIEW'),
    ('COMPUTATIONAL PROBLEM STATEMENT', 'III. COMPUTATIONAL PROBLEM STATEMENT'),
    ('PROPOSED SOLUTION', 'IV. PROPOSED SOLUTION'),
    ('METHODOLOGY AND MULTI-MODEL ML ARCHITECTURE', 'V. METHODOLOGY AND MULTI-MODEL ML ARCHITECTURE'),
    ('DATASET DESCRIPTION AND TEMPORAL DATA DISCIPLINE', 'VI. DATASET DESCRIPTION AND TEMPORAL DATA DISCIPLINE'),
    ('EXPLORATORY DATA ANALYSIS', 'VII. EXPLORATORY DATA ANALYSIS'),
    ('MODELS', 'VIII. MODELS'),
    ('MODEL EVALUATION', 'IX. MODEL EVALUATION'),
    ('MODEL SELECTION', 'X. MODEL SELECTION'),
    ('DEPLOYMENT', 'XI. DEPLOYMENT'),
    ('DISCUSSION', 'XII. DISCUSSION'),
    ('LIMITATIONS AND CONCLUSION', 'XIII. LIMITATIONS AND CONCLUSION'),
    ('REFERENCES', 'XIV. REFERENCES'),
]

mapping = dict(heading_order)

for p in doc.paragraphs:
    txt = (p.text or '').strip()

    if txt in mapping:
        p.text = mapping[txt]
        for run in p.runs:
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        continue

    if txt.startswith('ABSTRACT '):
        if not txt.startswith('ABSTRACT - '):
            p.text = txt.replace('ABSTRACT ', 'ABSTRACT - ', 1)
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        continue

    if txt.startswith('INDEX TERMS'):
        p.text = txt.replace('INDEX TERMS', 'INDEX TERMS -', 1) if not txt.startswith('INDEX TERMS -') else txt
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        p.paragraph_format.space_after = Pt(8)
        continue

    if txt.startswith(('A. ', 'B. ', 'C. ', 'D. ', 'E. ')):
        for run in p.runs:
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
    else:
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

for i in range(min(5, len(doc.paragraphs))):
    p = doc.paragraphs[i]
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    if i == 0:
        for run in p.runs:
            run.bold = True
    p.paragraph_format.space_after = Pt(2)

doc.save(path)
print('Applied section numbering and formal formatting updates.')
