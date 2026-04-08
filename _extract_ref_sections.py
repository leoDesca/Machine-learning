from docx import Document

doc = Document('REFINED CONCEPT PAPER 6.docx')
# print selected ranges around major sections
for i,p in enumerate(doc.paragraphs):
    t=(p.text or '').strip()
    if t in {'MODELS','MODEL EVALUATION','MODEL SELECTION','DEPLOYMENT'} or 'INTRODUCTION'==t or 'METHODOLOGY' in t:
        start=max(0,i-3); end=min(len(doc.paragraphs), i+30)
        print('\n' + '='*20 + f' around {i}: {t} ' + '='*20)
        for j in range(start,end):
            txt=(doc.paragraphs[j].text or '').strip()
            if txt:
                print(f'[{j:03}] {txt}')
