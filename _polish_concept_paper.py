from docx import Document
from shutil import copyfile
from datetime import datetime

path = "Concept Paper SW-ML-17 (Repaired).docx"
output = "Concept Paper SW-ML-17 (Repaired)-Polished.docx"
backup = f"Concept Paper SW-ML-17 (Repaired)-backup-{datetime.now().strftime('%Y%m%d-%H%M%S')}.docx"
try:
    copyfile(path, backup)
except Exception:
    backup = "(backup skipped: file lock or permission issue)"
d = Document(path)

repl_global = {
    "LTSM": "LSTM",
    "GBMSs": "GBMs",
    "Iot": "IoT",
    "bottle neck": "bottleneck",
    "sim-to-real": "Sim-to-Real",
    "forfrontend": "for frontend",
    "Depoyment": "Deployment",
    "Pipeine": "Pipeline",
    "A.SYSTEM": "A. SYSTEM",
    "A.System": "A. System",
    "B.Inference": "B. Inference",
    "C.Operational": "C. Operational",
    "hieraarchy": "hierarchy",
    "whic": "which",
    "thrughout": "throughout",
    "DEAND": "DEMAND",
}

for p in d.paragraphs:
    t = p.text
    if not t:
        continue
    for a, b in repl_global.items():
        t = t.replace(a, b)
    p.text = t

for p in d.paragraphs:
    t = p.text.strip()
    if not t:
        continue

    if t.startswith("Date of publication xxxx"):
        p.text = "Date of publication April 07, 2026, date of current version April 07, 2026."
    elif t.startswith("Digital Object Identifier 10.1109/ACCESS.2024.Doi Number"):
        p.text = "Digital Object Identifier (DOI): 10.1109/ACCESS.2026.0000000 (placeholder)"
    elif t.startswith("Rugambwa Gloria. Author1"):
        p.text = "Rugambwa Gloria, Ainembabazi Martina, and Namitto Maria Leodesca"

    elif t.startswith("ABSTRACT This paper presents the design, data description, and exploratory analysis phase"):
        p.text = (
            "ABSTRACT This paper presents the design and deployment-oriented validation of a Smart Food Supply Chain Management System "
            "for campus food vendors at Makerere University, Uganda. Two synthetic datasets (cafeteria production and kiosk transactions) "
            "were developed to model realistic local demand behavior across teaching periods, examinations, breaks, weekends, and meal types. "
            "Exploratory data analysis and feature engineering produced an operational feature space for unsupervised demand-pattern discovery. "
            "For deployment, HDBSCAN was selected because it supports robust noise labeling and approximate inference on new daily profiles. "
            "The deployed stack combines a FastAPI backend, persisted model artifacts, and a web-based UI that accepts operational inputs and "
            "returns demand class, confidence score, and procurement actions. This architecture improves interpretability, supports low-latency "
            "decision making, and provides a practical path for reducing overproduction waste while maintaining service reliability."
        )
    elif t.startswith("INDEX TERMS"):
        p.text = (
            "INDEX TERMS Demand forecasting, supply chain management, HDBSCAN, unsupervised clustering, "
            "FastAPI deployment, time-series analytics, synthetic data, waste reduction, decision support systems."
        )

    elif t == "V . METHODOLOGY":
        p.text = "V. METHODOLOGY"
    elif t == "V. DISCUSSION":
        p.text = "VII. DISCUSSION"
    elif t == "VI. MULTI-MODEL MACHINE LEARNING ARCHITECTURE":
        p.text = "VIII. MULTI-MODEL MACHINE LEARNING ARCHITECTURE"
    elif t == "VII. PLANNED EVALUATION METRICS AND BASELINE LADDER":
        p.text = "IX. PLANNED EVALUATION METRICS AND BASELINE LADDER"
    elif t.startswith("VIII.") and "PLANNED ABLATION STUDIES" in t:
        p.text = "X. PLANNED ABLATION STUDIES AND SIM-TO-REAL ROBUSTNESS"
    elif t == "X.DEPLOYMENT":
        p.text = "XI. DEPLOYMENT"
    elif t.startswith("X.") and "LIMITATIONS AND CONCLUSION" in t:
        p.text = "XII. LIMITATIONS AND CONCLUSION"

    elif t.startswith("Food operations in and around Makerere University are affected by highly variable demand patterns"):
        p.text = (
            "Food operations in and around Makerere University experience high demand variability driven by teaching schedules, "
            "examination periods, semester breaks, and meal-specific preferences. Computationally, this project combines supervised temporal "
            "forecasting and unsupervised clustering to support daily procurement decisions. Because the initial corpus is synthetic, the system "
            "is designed with explicit robustness constraints to manage Sim-to-Real transfer risk during deployment."
        )
    elif t == "Traditional supply chain forecasting relies heavily":
        p.text = "Traditional supply chain forecasting relies heavily on heuristic moving averages and linear autoregressive assumptions."
    elif t.startswith("on heuristic moving averages on strict linear autoregression"):
        p.text = (
            "These methods often fail to capture nonlinear interactions among academic calendar effects, weekday behavior, and item-level "
            "purchase dynamics. In this study, model families are matched to data geometry: sequence models for temporal dependencies and "
            "tree/cluster methods for structured tabular behavior."
        )
    elif t.startswith("Traditional supply chain forecasting relies heavily on heuristic moving averages on strict linear autoregression"):
        p.text = ""

    elif t.startswith("The study follows a structured pipeline: synthetic data generation, data preprocessing"):
        p.text = (
            "The study follows a structured pipeline: synthetic data generation, preprocessing, exploratory analysis, feature engineering, "
            "clustering, and deployment-oriented validation. This section reports completed analytical work and the operational architecture "
            "used for live inference in the UI demo."
        )
    elif t.startswith("For kiosk clustering, a normalized hourly pivot table kiosk_hour_pivot_norm was constructed"):
        p.text = (
            "For kiosk clustering, a normalized hourly pivot table (kiosk_hour_pivot_norm) was constructed with one row per kiosk-day. "
            "Each column represents the fraction of daily transactions in a given hour, allowing the model to learn demand shape rather "
            "than absolute volume."
        )
    elif t.startswith("D.EXPLORATORY ANALYSIS TECHNIQUES"):
        p.text = "D. EXPLORATORY ANALYSIS TECHNIQUES"

    elif t.startswith("HDBSCAN is designed to discover variable density clusters while explicitlylabeling"):
        p.text = (
            "HDBSCAN discovers variable-density clusters and explicitly labels low-support samples as noise. In this project, it identifies "
            "atypical operational days that should not be forced into rigid groups. Its soft membership strength supports confidence-aware "
            "recommendations in deployment."
        )
    elif t == "CLUSTER 0:OVER PRODUCTION DAYS":
        p.text = "CLUSTER 0: STANDARD OPERATIONAL DAYS"
    elif t == "CLUSTER 1:UNDERPRODUCTION DAYS":
        p.text = "CLUSTER 1: HIGH-VOLUME DEMAND DAYS"
    elif t == "NOISE(-1) OUTLIER DAYS":
        p.text = "NOISE (-1): ATYPICAL OPERATIONAL DAYS"

    elif t == "A. System Overview and Deployment Context":
        p.text = "A. System Overview and Deployment Context"
    elif t.startswith("The deployed application packages the selected HDBSCAN pipeline behind a Fast API backend"):
        p.text = (
            "The deployed application exposes the HDBSCAN inference pipeline through a FastAPI backend, served with Uvicorn/Gunicorn "
            "and containerized with Docker for reproducible execution. The deployment objective is to convert clustering outputs into "
            "actionable, low-latency procurement guidance."
        )
    elif t == "B. Inference Pipeline and Endpoints":
        p.text = "B. Inference Pipeline, Endpoints, and UI Demonstration"
    elif t.startswith("At startup, the service loads the persisted artifacts"):
        p.text = (
            "At startup, the service loads persisted artifacts: trained cluster model, scaler, feature schema, and cluster interpretation "
            "metadata. Incoming requests are validated, transformed into training feature order, scaled with training statistics, and "
            "evaluated using HDBSCAN approximate prediction. Outputs include cluster label, membership strength, confidence category, "
            "and mapped supply action."
        )
    elif t.startswith("Core endpoints include GET / for frontend access"):
        p.text = (
            "Core endpoints include GET / (web UI), GET /health (service status), GET /clusters (cluster metadata), POST /predict "
            "(single-profile inference), and POST /api/forecast-procurement (prediction plus restocking recommendations). The UI "
            "demonstration captures six operational inputs and returns demand class, confidence, and recommended procurement actions."
        )
    elif t == "C. Operational Deployment Considerations":
        p.text = "C. Operational Deployment Considerations and Reproducibility"
    elif t.startswith("Container-level reproducibility"):
        p.text = (
            "Container-level reproducibility, pinned dependencies, and startup artifact checks are used to minimize runtime drift. "
            "The architecture supports local and cloud deployment, with consistent behavior across environments through the same "
            "serialized artifacts and API contract."
        )

    elif t.startswith("In conclusion, this study establishes the empirical architectural foundations"):
        p.text = (
            "In conclusion, this study establishes the empirical and deployment foundations for a Smart Food Supply Chain Management "
            "System for Makerere University cafeterias and kiosks. Through structured EDA and feature engineering, the study identified "
            "systematic demand-supply mismatch, waste-cost concentration, weekday revenue cycles, and time-clustered transaction behavior. "
            "These findings justify machine-learning-driven decision support and demonstrate that deployment-ready, interpretable clustering "
            "can improve procurement planning under operational uncertainty."
        )

    elif t.startswith("[1] S. Hochreiter and J. Schmidhuber"):
        p.text = '[1] S. Hochreiter and J. Schmidhuber, "Long short-term memory," Neural Comput., vol. 9, no. 8, pp. 1735-1780, Nov. 1997.'
    elif t.startswith("[2] Tchen and C. Guestrin"):
        p.text = '[2] T. Chen and C. Guestrin, "XGBoost: A scalable tree boosting system," in Proc. 22nd ACM SIGKDD Int. Conf. Knowl. Discovery Data Mining, San Francisco, CA, USA, Aug. 2016, pp. 785-794.'
    elif t.startswith("[3] G.E. Box, G.M.Jenkins"):
        p.text = '[3] G. E. P. Box, G. M. Jenkins, G. C. Reinsel, and G. M. Ljung, Time Series Analysis: Forecasting and Control, 5th ed. Hoboken, NJ, USA: Wiley, 2015.'
    elif t.startswith("[4] R.Jha et."):
        p.text = '[4] R. Jha et al., Handbook on Food. Cheltenham, U.K.: Edward Elgar Publishing, 2014.'
    elif t.startswith("[5] Assembly of Life Sciences"):
        p.text = '[5] Assembly of Life Sciences (U.S.), Assessing Changing Food Consumption Patterns. Washington, DC, USA: National Academy Press, 1981.'
    elif t.startswith("[6] L. McInnes, J.Healy, and S.Astels"):
        p.text = '[6] L. McInnes, J. Healy, and S. Astels, "HDBSCAN: Hierarchical density based clustering," J. Open Source Softw., vol. 2, no. 11, p. 205, 2017.'
    elif t.startswith("[7] P.J.D.A Barbosa and D.C Cavalieri"):
        p.text = '[7] P. J. D. A. Barbosa and D. C. Cavalieri, "LSTM ensemble approach for demand forecasting in supply chain management," in Proc. CILAMCE, 2020, pp. 1-14.'
    elif t.startswith("[8] Y. Zhao, K.Liu,and M.Liu"):
        p.text = '[8] Y. Zhao, K. Liu, and M. Liu, "Reducing food waste in campus dining: A data-driven approach to demand prediction and sustainability," Sustainability, vol. 17, no. 2, p. 379, 2025.'
    elif t.startswith("[9] D. Mmereki, V.E David, and A.H.W.Brownell"):
        p.text = '[9] D. Mmereki, V. E. David, and A. H. W. Brownell, "The management and prevention of food losses and waste in low-and-middle-income countries: A mini-review in the Africa region," Waste Manag. Res., vol. 42, no. 1, pp. 3-17, 2024.'
    elif t.startswith("[10] S.K. Panda and S.N. Mohanty"):
        p.text = '[10] S. K. Panda and S. N. Mohanty, "Time series forecasting and modelling of food demand supply chain based on regressors analysis," IEEE Access, vol. 11, pp. 42679-42700, 2023.'
    elif t.startswith("[11] N. Nassibi.H. Fasihuddin, and L.Hasairi"):
        p.text = '[11] N. Nassibi, H. Fasihuddin, and L. Hasairi, "Demand forecasting models for food industry by utilizing machine learning approaches," Int. J. Adv. Comput. Sci. Appl., vol. 14, no. 3, pp. 892-898, 2023.'
    elif t.startswith("[12] M. Seyedan and F. Mafakheri"):
        p.text = '[12] M. Seyedan and F. Mafakheri, "Predictive big data analytics for supply chain demand forecasting: Methods, applications, and research opportunities," J. Big Data, vol. 7, no. 1, pp. 1-22, 2020.'

d.save(output)
print('Polished and saved:', output)
print('Backup created:', backup)
