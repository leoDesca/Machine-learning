from docx import Document
from docx.shared import Pt

out = Document()

# Basic style tuning for readable academic format
style = out.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

# Title block
out.add_paragraph('A Multi-Model Machine Learning Architecture for Food Supply Chain Segmentation and Optimization in Makerere University Food Operations.')
out.add_paragraph('SW-ML-17')
out.add_paragraph('RUGAMBWA GLORIA, AINEMBABAZI MARTINA and NAMITTO MARIA LEODESCA')
out.add_paragraph('Department of Computer Science, COCIS, Makerere University, Kampala, Uganda')
out.add_paragraph('Supervisor: Ggaliwango Marvin')

out.add_paragraph(
    'ABSTRACT This concept paper presents a detailed machine learning architecture for operational segmentation and decision support in Makerere University food supply systems. '
    'The study is based on two synthetic but context-grounded datasets: a cafeteria production dataset and a kiosk transaction dataset that encode realistic Ugandan consumption behavior, '
    'academic calendar shifts, and campus demand cycles. The analysis pipeline combines structured preprocessing, exploratory data analysis, feature engineering, and model assessment under '
    'a multi-model unsupervised learning design. Four clustering models are developed and compared: HDBSCAN, Spectral Clustering, Agglomerative Clustering, and BIRCH. '
    'Each model contributes a distinct inductive bias for density detection, graph-based partitioning, hierarchical grouping, and scalable incremental clustering, respectively. '
    'The deployment layer operationalizes the selected HDBSCAN model through a FastAPI and Docker architecture, exposing health, cluster metadata, and prediction endpoints for live procurement guidance. '
    'The resulting system is designed to reduce overproduction waste, improve stock planning, and increase vendor profitability while maintaining interpretable and maintainable ML behavior.'
)
out.add_paragraph(
    'INDEX TERMS Food supply chain management, synthetic data, unsupervised clustering, HDBSCAN, Spectral Clustering, Agglomerative Clustering, BIRCH, waste reduction, deployment architecture.'
)

out.add_paragraph('INTRODUCTION')
out.add_paragraph(
    'Food operations in and around Makerere University are affected by highly variable demand patterns driven by teaching schedules, examination periods, semester breaks, and meal-specific preferences. '
    'In this context, uniform procurement policies often create costly mismatch between prepared portions and actual sales, causing avoidable waste and unstable profit margins for campus vendors.'
)
out.add_paragraph(
    'This project addresses that challenge as an operational segmentation problem instead of a pure forecasting problem. Rather than predicting a single next value, the system groups similar daily '
    'operational states into actionable clusters that represent low-risk, medium-risk, and high-risk supply conditions. This enables vendors to map each incoming day profile to a data-driven procurement posture.'
)
out.add_paragraph(
    'The paper is organized to mirror an engineering deployment lifecycle: computational problem definition, dataset discipline, methodology, model-by-model design, evaluation, selection, and deployment. '
    'This structure allows each architectural decision to be justified technically and traced to practical implementation.'
)

out.add_paragraph('BACKGROUND AND LITERATURE REVIEW')
out.add_paragraph(
    'Traditional supply chain studies in institutional food services frequently prioritize forecasting-based methods such as moving averages, ARIMA variants, and regression pipelines. '
    'While useful in stable environments, these methods can underperform when behavior is multi-modal and regime-dependent, as in campus operations where the same meal can show very different demand under different academic periods.'
)
out.add_paragraph(
    'Clustering-based approaches provide an alternative by discovering latent operational regimes directly from feature geometry. Density-based methods handle outliers and irregular cluster shapes; '
    'graph-based methods capture non-linear boundaries; hierarchical methods provide interpretable coarse-to-fine segmentation; and incremental clustering methods support scalable updates.'
)
out.add_paragraph(
    'The present architecture integrates these complementary families in one comparative framework, then deploys the model that best balances separation quality, interpretability, and production feasibility.'
)

out.add_paragraph('COMPUTATIONAL PROBLEM STATEMENT')
out.add_paragraph(
    'The core computational task is to learn robust operational clusters from temporally indexed tabular data generated under sim-to-real constraints. '
    'The absence of fully instrumented real-time sensing required synthetic generation of training corpora informed by local food behavior. This introduces domain-shift risk when models are applied to live settings.'
)
out.add_paragraph(
    'Optimization risk appears when synthetic regularities are overlearned and then transferred directly to noisy real conditions. In practice, this can manifest as incorrect cluster assignment and suboptimal recommendations '
    'that trigger overstocking, underproduction, or unnecessary emergency procurement.'
)
out.add_paragraph(
    'To mitigate this risk, the architecture is designed for continuous adaptation: synthetic data is treated as structured initialization, while deployment is built to accept rolling updates and local recalibration as empirical data accumulates.'
)

out.add_paragraph('PROPOSED SOLUTION')
out.add_paragraph(
    'The proposed solution is a multi-model clustering decision system with a single selected deployment model. The system constructs engineered operational features for cafeteria-day and kiosk-day units, '
    'fits multiple clustering algorithms with different assumptions, compares quality and business utility, and exposes the selected model as an API service for day-level inference.'
)
out.add_paragraph(
    'The deployed prediction output is not just a cluster ID. It includes membership strength and a mapped operational action label to support practical decision-making, for example conservative production, '
    'balanced preparation, or demand-responsive increase in output.'
)

out.add_paragraph('METHODOLOGY AND MULTI-MODEL ML ARCHITECTURE')
out.add_paragraph('A. Data Preprocessing')
out.add_paragraph(
    'Both datasets are ingested from CSV, validated for missingness and duplicates, cleaned for currency formatting, and standardized for datetime consistency. '
    'Academic period and weekday are encoded as ordered categorical variables to preserve calendar semantics in downstream feature generation.'
)
out.add_paragraph('B. Feature Engineering Pipeline')
out.add_paragraph(
    'Cafeteria features include production volume, sell-out indicators, waste percentage, gross margin proxies, meal diversity entropy, weekday cyclical encodings, and short rolling windows. '
    'Kiosk features include normalized hourly transaction distributions to represent demand shape independent of absolute volume.'
)
out.add_paragraph('C. Exploratory Analysis Layer')
out.add_paragraph(
    'EDA includes mismatch analysis, waste distribution profiling, meal-day heatmaps, hourly demand concentration, and comparative trend diagnostics across academic periods. '
    'These analyses inform which features best separate operational regimes before clustering.'
)
out.add_paragraph('D. Multi-Model Training and Comparison')
out.add_paragraph(
    'Four clustering models are trained over standardized feature spaces and assessed using internal validity and operational interpretability. '
    'Comparative diagnostics include silhouette behavior, Davies-Bouldin compactness, Calinski-Harabasz separation, cluster population stability, and practical meaning of cluster-specific action policies.'
)

out.add_paragraph('DATASET DESCRIPTION AND TEMPORAL DATA DISCIPLINE')
out.add_paragraph('Cafeteria Dataset')
out.add_paragraph(
    'The cafeteria dataset represents 14 preparation units over approximately two academic years and includes six core meal categories. '
    'Features capture prepared portions, sold portions, waste portions, waste rates, pricing, revenue, estimated ingredient costs, and gross profit signals with day and academic-period context.'
)
out.add_paragraph('Kiosk Dataset')
out.add_paragraph(
    'The kiosk dataset contains transaction-level records for seven kiosks and nine common campus items, including timestamped purchase behavior and payment channels. '
    'Hourly demand weighting follows realistic lecture-driven movement patterns with expected peaks around key transition periods in the academic day.'
)
out.add_paragraph('Temporal Discipline')
out.add_paragraph(
    'All transformations preserve chronological semantics. Rolling features are computed in causal order, and validation logic is designed to avoid temporal leakage. '
    'This is essential for reliable transfer of learned cluster structure into live operational settings.'
)

out.add_paragraph('EXPLORATORY DATA ANALYSIS')
out.add_paragraph(
    'Exploratory analysis confirms recurring demand-supply mismatch across multiple meal categories, with identifiable high-waste regimes and period-specific demand swings. '
    'Cafeteria profiles exhibit non-uniform waste burdens, while kiosk transaction shapes reveal reproducible temporal concentration windows that influence replenishment pressure.'
)
out.add_paragraph(
    'The EDA stage also supports feature pruning by identifying redundancy and emphasizing variables with strong regime-discriminative behavior. '
    'This ensures clustering is driven by operationally meaningful structure rather than noise-dominated dimensions.'
)

out.add_paragraph('MODELS')
out.add_paragraph('A. Model 1: HDBSCAN (Density-Based Clustering)')
out.add_paragraph(
    'HDBSCAN is designed to discover variable-density clusters while explicitly labeling low-support samples as noise. '
    'In this project, HDBSCAN is particularly useful for identifying atypical cafeteria-day patterns that should not be forced into rigid clusters. '
    'Its soft assignment capabilities support confidence-aware recommendation behavior in deployment.'
)
out.add_paragraph('B. Model 2: Spectral Clustering (Graph-Based Partitioning)')
out.add_paragraph(
    'Spectral Clustering captures non-linear operational boundaries by embedding samples in a graph Laplacian space before partitioning. '
    'This model is effective when demand regimes are connected through manifold-like structure rather than separable by simple convex assumptions.'
)
out.add_paragraph('C. Model 3: Agglomerative Clustering (Hierarchical Grouping)')
out.add_paragraph(
    'Agglomerative Clustering builds a hierarchy of mergers from fine-grained local similarities to broader operational regimes. '
    'Its output supports managerial interpretation by allowing cluster inspection at multiple granularity levels, which is useful for policy communication to non-technical stakeholders.'
)
out.add_paragraph('D. Model 4: BIRCH (Balanced Iterative Reducing and Clustering using Hierarchies)')
out.add_paragraph(
    'BIRCH compresses large sample streams into clustering features and is computationally efficient for incremental scenarios. '
    'Within this study, BIRCH provides a scalability baseline for future extension when real-time records increase beyond synthetic corpus size.'
)
out.add_paragraph('E. Model Integration Perspective')
out.add_paragraph(
    'The four-model design is intentionally comparative rather than redundant. Each model exposes a different geometric assumption and operational trade-off. '
    'The final deployment choice is made after balancing internal metrics, cluster stability, interpretability, and API-serving practicality.'
)

out.add_paragraph('MODEL EVALUATION')
out.add_paragraph('A. Cross-Model Evaluation Criteria')
out.add_paragraph(
    'Evaluation combines internal clustering validity with business relevance. Quantitative criteria include silhouette score, Davies-Bouldin index, and Calinski-Harabasz index. '
    'Operational criteria include actionable cluster narratives, noise handling behavior, and consistency under temporal slices.'
)
out.add_paragraph('B. Comparative Findings')
out.add_paragraph(
    'HDBSCAN achieved the best balance between separation quality and operational robustness, with strong handling of irregular demand regimes and uncertain points. '
    'Spectral Clustering captured non-linear partitions effectively but required careful affinity tuning. Agglomerative Clustering provided interpretable hierarchy but was more sensitive to linkage design. '
    'BIRCH showed efficient scaling and reasonable partition quality, making it suitable for growth-oriented scenarios.'
)
out.add_paragraph(
    'Across runs, the deployed HDBSCAN variant produced stable cluster-action mappings and practical confidence behavior, making it the strongest candidate for production integration.'
)

out.add_paragraph('MODEL SELECTION')
out.add_paragraph(
    'HDBSCAN was selected for deployment because it simultaneously satisfies four requirements: robust cluster discovery under variable density, explicit treatment of anomalous points as noise, '
    'confidence-aware assignment through approximate prediction, and straightforward integration into a lightweight API service. '
    'This combination outperformed alternatives in the context of real-time operational advisory requirements.'
)

out.add_paragraph('DEPLOYMENT')
out.add_paragraph('A. System Overview and Deployment Context')
out.add_paragraph(
    'The deployed application packages the selected HDBSCAN pipeline behind a FastAPI backend, served with Gunicorn and Uvicorn workers, and containerized with Docker for reproducible cloud execution. '
    'The deployment objective is to transform research clustering outputs into actionable, low-latency vendor guidance.'
)
out.add_paragraph('B. Inference Pipeline and Endpoints')
out.add_paragraph(
    'At startup, the service loads the persisted artifacts: clustering model, scaler, feature schema, and cluster interpretation metadata. '
    'Incoming requests are validated, transformed into model feature order, scaled consistently with training statistics, and evaluated through HDBSCAN approximate prediction. '
    'Returned outputs include cluster label, membership strength, and mapped supply action for decision support.'
)
out.add_paragraph(
    'Core endpoints include GET / for frontend access, GET /health for service and artifact status, GET /clusters for cluster interpretation lookup, '
    'and POST /predict for operational recommendation on new daily profiles.'
)
out.add_paragraph('C. Operational Deployment Considerations')
out.add_paragraph(
    'Container-level reproducibility, explicit dependency locking, and startup artifact checks are used to minimize runtime drift. '
    'The architecture is designed for cloud deployment through Docker-capable platforms, with blueprint-based rollout available through Render configuration.'
)

out.add_paragraph('DISCUSSION')
out.add_paragraph(
    'This work demonstrates that a clustering-first strategy is appropriate for institutional food operations where behavior is regime-driven and not strictly deterministic over time. '
    'By separating model families and evaluating them under common criteria, the study provides both technical clarity and managerial interpretability. '
    'The deployment bridge further validates that the architecture is not only analytical but operationally usable.'
)

out.add_paragraph('LIMITATIONS AND CONCLUSION')
out.add_paragraph(
    'The primary limitation is reliance on synthetic training distributions, which may not capture all real-world anomalies, policy shocks, or exogenous disruptions. '
    'Future work should incorporate live telemetry, periodic recalibration, and semi-supervised adaptation to reduce sim-to-real uncertainty.'
)
out.add_paragraph(
    'In conclusion, the project establishes a complete multi-model ML framework for Makerere food supply segmentation, from data construction and EDA through model comparison, selection, and deployment. '
    'The HDBSCAN-centered production design provides a practical path toward waste reduction, more stable procurement, and evidence-based campus food operations management.'
)

out.add_paragraph('REFERENCES')
out.add_paragraph('[1] Ester, M., Kriegel, H.-P., Sander, J., & Xu, X. (1996). A density-based algorithm for discovering clusters in large spatial databases with noise. Proceedings of KDD.')
out.add_paragraph('[2] Campello, R. J. G. B., Moulavi, D., & Sander, J. (2013). Density-based clustering based on hierarchical density estimates. PAKDD.')
out.add_paragraph('[3] Ng, A. Y., Jordan, M. I., & Weiss, Y. (2002). On spectral clustering: Analysis and an algorithm. NIPS.')
out.add_paragraph('[4] Rokach, L., & Maimon, O. (2005). Clustering methods. In Data Mining and Knowledge Discovery Handbook.')
out.add_paragraph('[5] Zhang, T., Ramakrishnan, R., & Livny, M. (1996). BIRCH: An efficient data clustering method for very large databases. SIGMOD.')
out.add_paragraph('[6] McInnes, L., Healy, J., & Astels, S. (2017). hdbscan: Hierarchical density based clustering. Journal of Open Source Software.')
out.add_paragraph('[7] Pedregosa, F., et al. (2011). Scikit-learn: Machine learning in Python. Journal of Machine Learning Research.')
out.add_paragraph('[8] FastAPI Documentation. https://fastapi.tiangolo.com')
out.add_paragraph('[9] Docker Documentation. https://docs.docker.com')

out.save('SW-ML-17 Concept Paper.docx')
print('Rewrote SW-ML-17 Concept Paper.docx with reference-style structure and detailed model/deployment sections.')
